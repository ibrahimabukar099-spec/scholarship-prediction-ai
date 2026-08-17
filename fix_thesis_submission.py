from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


INPUT = Path(r"C:\Users\Ahmed\Desktop\SEPS1_5_formatted.docx")
OUTPUT = Path.cwd() / "SEPS1_5_corrected_submission.docx"


def set_para_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def clone_paragraph_before_table(table, source_paragraph, text):
    new_p = deepcopy(source_paragraph._p)
    table._tbl.addprevious(new_p)
    paragraph = source_paragraph._parent.paragraphs[-1]
    # The copied paragraph is not reliably exposed as the last paragraph in python-docx,
    # so create a wrapper by scanning the XML-backed paragraphs after insertion.
    for p in source_paragraph._parent.paragraphs:
        if p._p is new_p:
            paragraph = p
            break
    set_para_text(paragraph, text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return paragraph


doc = Document(INPUT)

# Chapter 1: add Somalia-specific evidence and remove the unsupported HDLN mention.
set_para_text(
    doc.paragraphs[3],
    "A scholarship is a form of financial assistance provided to students, often by governments, non-governmental organizations, or academic institutions, to make higher education more accessible and affordable. As the cost of attending university continues to rise globally, scholarships have become a critical mechanism for ensuring social equity in the education sector. In the specific context of Somalia, higher education institutions and education administrators operate within a sector that is still strengthening governance, data systems, access, equity, and ICT-supported planning. Somalia's National Education Sector Strategic Plan 2022-2026 identifies education governance, higher education, ICT, and equity as priority areas, while the World Bank Group's education management information system assessment highlights the importance of stronger education data for planning and decision-making in Somalia (Federal Government of Somalia, Ministry of Education, Culture and Higher Education, 2022; World Bank Group, 2018). UNICEF Somalia (2025) also reports that conflict, climate shocks, and displacement continue to affect children and families, reinforcing the need for transparent and context-aware education support mechanisms.",
)

set_para_text(
    doc.paragraphs[37],
    "Chapter Two, Literature Review, provides an in-depth analysis of existing scholarly works, theories, and models related to machine learning and scholarship management. This chapter explores relevant approaches such as Naive Bayes, Random Forest, Logistic Regression, SVM, and XGBoost as discussed in recent indexed papers. It also reviews previous studies on socio-economic and academic factors that influence eligibility, providing a theoretical framework for the current system.",
)

set_para_text(
    doc.paragraphs[75],
    "Many existing scholarship prediction studies are based on data from institutions outside Somalia and East Africa. This creates a contextual limitation because scholarship eligibility factors differ across regions. For example, displacement, orphan status, regional inequality, and family income patterns are particularly important in the Somali context. UNHCR describes Somalia as a country affected by conflict, insecurity, climate-related crises, and large-scale internal displacement, while UNICEF Somalia (2025) links recent climate shocks and displacement to wider humanitarian pressure on children and families (UNHCR, n.d.; UNICEF Somalia, 2025). If a model is trained or designed using assumptions from another region, it may fail to represent the realities faced by Somali students.",
)

set_para_text(
    doc.paragraphs[76],
    "To address this gap, the current study uses a structured project dataset named scholarship_dataset.csv. The dataset contains 2,000 applicant records with nine active prediction features: GPA, family income, orphan status, displaced status, region, high school type, verification status, gender, and faculty, along with the target eligibility status. Since direct institutional scholarship data was not available for open use, the dataset is treated as a synthetic project dataset calibrated to reflect Somali socio-economic and educational eligibility factors identified in the literature and in Somalia-specific policy and humanitarian sources (Federal Government of Somalia, Ministry of Education, Culture and Higher Education, 2022; UNICEF Somalia, 2025; UNHCR, n.d.). This makes the dataset suitable for prototype development and methodological demonstration, while the study acknowledges that future deployment should be validated using real institutional data.",
)

# Remove duplicated Fajardo paragraph.
if doc.paragraphs[61].text.strip() == doc.paragraphs[62].text.strip():
    doc.paragraphs[62]._element.getparent().remove(doc.paragraphs[62]._element)

# Chapter 3: make NFRs grammatically correct and honest about untested performance/reliability.
for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    if text.startswith("NFR1 Performance:"):
        set_para_text(
            paragraph,
            "NFR1 Performance: The system is intended to return prediction results within 3 seconds for individual prototype user requests. This is a theoretical performance requirement rather than a load-tested result; formal concurrency and load testing were outside the scope of this prototype evaluation.\nNFR2 Usability: The interface shall remain readable and usable on screen widths from 375px to 1440px.",
        )
    elif text.startswith("NFR4 Reliability:"):
        set_para_text(
            paragraph,
            "NFR4 Reliability: The system architecture is intended to store submitted application records consistently during normal prototype use. Because formal concurrency and load testing were not conducted, reliability under moderate concurrent submissions remains a requirement for future validation rather than a confirmed test result.\nNFR5 Explainability: Each saved prediction shall provide an explanation text and, when SHAP calculation is available, a SHAP chart showing feature contribution values.",
        )

# Chapter 4: add missing Table 4.1 caption before the functional testing table.
caption_exists = any("Table 4.1: Black-Box Functional Test Cases" in p.text for p in doc.paragraphs)
if not caption_exists:
    # Table 4 is the black-box functional test cases table.
    template_caption = next((p for p in doc.paragraphs if p.text.strip().startswith("Table 4.2:")), doc.paragraphs[0])
    clone_paragraph_before_table(doc.tables[4], template_caption, "Table 4.1: Black-Box Functional Test Cases")

# Chapter 5: qualify Objective 4 achievement because evaluation used a synthetic dataset.
objective_table = doc.tables[10]
objective_table.cell(4, 1).text = "Substantially Achieved"
objective_table.cell(4, 2).text = (
    "The four models were evaluated using accuracy, precision, recall, F1-score, ROC-AUC, confusion matrix, "
    "and 5-fold cross-validation on the project dataset. The selected XGBoost model achieved 97.50% accuracy, "
    "93.85% precision, 98.39% recall, 96.06% F1-score, 99.77% ROC-AUC, and confusion matrix [[268, 8], [2, 122]]. "
    "However, because the dataset is synthetic and no real institutional validation was conducted, the objective is "
    "reported as substantially achieved at prototype level rather than fully validated for deployment."
)

# References: add missing Somalia-specific APA entries before APPENDIX.
reference_texts = [
    "Federal Government of Somalia, Ministry of Education, Culture and Higher Education. (2022). National education sector strategic plan 2022-2026. https://www.iicba.unesco.org/en/africa-education-knowledge-platform/national-education-sector-strategic-plan-2022-2026",
    "UNHCR. (n.d.). Somalia. Retrieved August 16, 2026, from https://www.unhcr.org/where-we-work/countries/somalia",
    "UNICEF Somalia. (2025). 2024 UNICEF Somalia annual report. https://www.unicef.org/somalia/reports/2024-som-annual-report",
    "World Bank Group. (2018). Somalia education programmatic technical assistance: Status of education management information system. https://documents.worldbank.org/en/publication/documents-reports/documentdetail/317341555399356772",
]
existing_all = "\n".join(p.text for p in doc.paragraphs)
appendix_para = next(p for p in doc.paragraphs if p.text.strip() == "APPENDIX")
for ref in reversed(reference_texts):
    if ref.split(". (")[0] not in existing_all:
        new_p = deepcopy(appendix_para._p)
        appendix_para._p.addprevious(new_p)
        for p in doc.paragraphs:
            if p._p is new_p:
                set_para_text(p, ref)
                break

doc.save(OUTPUT)
print(OUTPUT)
